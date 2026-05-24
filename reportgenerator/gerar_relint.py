"""
Geração automática de Relatórios de Inteligência de Área (RELINTs)
a partir dos dados estruturados do CoPatrulha.

Uso:
    uv run python gerar_relint.py              # gera todos os 8 relatórios
    uv run python gerar_relint.py "Jardim"     # área cujo nome contenha "Jardim"

Saída: relints_gerados/RI_<num>_2026_<slug>.docx
"""

import json
import pathlib
import sys
from collections import Counter

from dotenv import load_dotenv
load_dotenv()

import geopandas as gpd
import pandas as pd
import polars as pl
from anthropic import Anthropic
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

CRS = "EPSG:4326"
RELINTS_DIR = pathlib.Path("relints")
OUTPUT_DIR = pathlib.Path("relints_gerados")

# Mapping: shapefile nome_subar → (ri_num, titulo uppercase, nome_area_fm em cameras)
AREA_CONFIG: dict[str, dict] = {
    "Rodoviária - Terminal Gentileza - Estação Leopoldina": {
        "ri_num": "010",
        "titulo": "RODOVIÁRIA – TERMINAL GENTILEZA – ESTAÇÃO LEOPOLDINA",
        "cam_nome": "Rodoviária - Terminal Gentileza - Estação Leopoldina",
    },
    "Metrô Botafogo - Rua São Clemente - Rua Voluntários da Pátria": {
        "ri_num": "011",
        "titulo": "METRÔ BOTAFOGO – RUA SÃO CLEMENTE – VOLUNTÁRIOS DA PÁTRIA",
        "cam_nome": "Metrô Botafogo - Rua São Clemente - Rua Voluntários da Pátria",
    },
    "Jardim de Alah": {
        "ri_num": "012",
        "titulo": "JARDIM DE ALAH",
        "cam_nome": "Jardim de Alah",
    },
    "Campo Grande: Estação de Trem - Calçadão": {
        "ri_num": "013",
        "titulo": "CAMPO GRANDE – ESTAÇÃO – CALÇADÃO",
        "cam_nome": "Campo Grande: Estação de Trem - Calçadão",
    },
    "Rio Sul": {
        "ri_num": "014",
        "titulo": "RIO SUL",
        "cam_nome": "Rua Lauro Müller – Avenida General Severiano – Avenida Venceslau Brás",
    },
    "Praia de Botafogo - Rua Marquês de Abrantes": {
        "ri_num": "015",
        "titulo": "PRAIA DE BOTAFOGO – MARQUÊS DE ABRANTES",
        "cam_nome": "Praia de Botafogo - Rua Marquês de Abrantes",
    },
    "Estações São Francisco Xavier - Afonso Pena": {
        "ri_num": "016",
        "titulo": "ESTAÇÕES SÃO FRANCISCO XAVIER – AFONSO PENA",
        "cam_nome": "Estações São Francisco Xavier - Afonso Pena",
    },
    "Presidente Vargas - Campo de Santana - Central do Brasil - Cinelândia": {
        "ri_num": "017",
        "titulo": "PRESIDENTE VARGAS – CAMPO DE SANTANA – CENTRAL – CINELÂNDIA",
        "cam_nome": "Presidente Vargas - Campo de Santana - Central do Brasil - Cinelândia",
    },
}

SYSTEM_PROMPT = """\
Você é analista de inteligência criminal do CoPatrulha Municipal do Rio de Janeiro.
Sua audiência é a alta gestão municipal — gestores que tomam decisões operacionais
em reuniões semanais. Escreva para quem lê rápido e precisa agir.

## PRINCÍPIOS DE ESCRITA
- Frase curta. Dado primeiro, contexto depois.
- Cite números e logradouros reais dos dados fornecidos.
- Elimine adjetivos decorativos ("intenso", "elevado", "extenso") — use cifras no lugar.
- Nada de repetição entre sub-áreas ou entre abertura e conclusão.
- Cada frase deve justificar uma decisão ou revelar um risco concreto.

## ESTRUTURA

**Introdução** (2–3 frases): o que torna esta área prioritária — números e localização.

**Cada sub-área**:
- Abertura: 2–4 frases. Vulnerabilidade principal + dado quantitativo + horário crítico.
- 5 bullets (sem •, adicionado pelo sistema):
  1. retenção de fluxo em horários de pico — [local + horário específico]
  2. áreas com baixa visibilidade — [trecho + causa]
  3. obstáculos urbanos dificultando vigilância — [obstáculo + impacto operacional]
  4. circulação intensa de motocicletas e bicicletas — [uso como rota de fuga]
  5. múltiplas rotas de dispersão após a prática criminosa — [vias — USE a rota do RELINT original quando fornecida]
- Dinâmica: 1 frase. Perfil do infrator + modus operandi + horário.

**Conclusão**:
- Abertura: 2–3 frases sintetizando o risco principal da área. Termina com "Observa-se necessidade de:"
- 5 recomendações (sem •): ação + responsável + local/horário específico.
- Fechamento: 1 frase sobre o padrão temporal dos delitos (quando e onde concentram).

## EXEMPLO

Introdução:
"A Rua São Clemente registra 821 ocorrências no período, com pico entre 18h e 22h.
Furtos de celular representam 68% dos delitos — praticados principalmente nas imediações do Metrô Botafogo e em trechos com arborização densa bloqueando a iluminação."

Sub-área (abertura):
"O acesso sul do Metrô Botafogo concentra 34% das ocorrências da área em 80m² de calçada.
O pico ocorre entre 17h e 19h, quando o fluxo de saída cria aglomeração contínua nas escadas.
Bancas e grades estreitam o corredor, eliminando a possibilidade de patrulhamento a pé em par."

Bullet 5: múltiplas rotas de dispersão após a prática criminosa — ligação rápida com Rua Humaitá e acesso ao Túnel Santa Bárbara

Dinâmica: "Indivíduos a pé e em moto atuam em duplas nas escadas do metrô, abordando vítimas no momento do pagamento da tarifa entre 17h e 19h."

Recomendação: "reforço do patrulhamento a pé no acesso sul do Metrô Botafogo — missão dirigida das 17h às 19h, dois agentes fixos na escada de saída"

## SAÍDA
Retorne EXCLUSIVAMENTE um JSON válido:
```json
{
  "intro": "...",
  "subareas": [
    {
      "abertura": "...",
      "fatores": ["retenção de fluxo...", "áreas com baixa visibilidade...", "obstáculos urbanos...", "circulação intensa...", "múltiplas rotas de dispersão..."],
      "dinamica": "..."
    },
    {...},
    {...}
  ],
  "conclusao": {
    "abertura": "...síntese... Observa-se necessidade de:",
    "recomendacoes": ["ação — responsável — local/horário", "...", "...", "...", "..."],
    "fechamento": "..."
  }
}
```
"""


# ── 1. Extração dos RELINTs existentes ──────────────────────────────────────

def extrair_relints() -> dict[str, dict]:
    """Parseia os 8 RELINTs existentes e extrai nomes de sub-áreas e rotas de fuga."""
    titulo_to_subar = {v["titulo"]: k for k, v in AREA_CONFIG.items()}
    resultado: dict[str, dict] = {}

    for fpath in sorted(RELINTS_DIR.glob("*.docx")):
        doc = Document(fpath)
        table = doc.tables[0]
        titulo = table.rows[1].cells[0].text.strip()

        nome_subar = titulo_to_subar.get(titulo)
        if not nome_subar:
            # Try normalizing dashes
            for t, ns in titulo_to_subar.items():
                if titulo.replace("—", "–") == t or titulo.replace("–", "—") == t:
                    nome_subar = ns
                    break
        if not nome_subar:
            print(f"  AVISO: sem mapeamento para '{titulo}' ({fpath.name})")
            continue

        subareas = []
        for row_idx in [3, 5, 7]:
            subarea_nome = table.rows[row_idx].cells[0].text.strip()
            content_cell = table.rows[row_idx + 1].cells[0]
            paras = [p.text for p in content_cell.paragraphs]

            # Bullet curto com "rotas de dispersão" = rota de fuga extraída
            escape = next(
                (p for p in paras if ("rotas de dispersão" in p.lower() or "dispersão" in p.lower()) and len(p) < 300),
                None,
            )
            subareas.append({
                "nome": subarea_nome,
                "escape_route": escape,
            })

        resultado[nome_subar] = {
            "subareas": subareas,
            "fonte": fpath.name,
        }

    return resultado


# ── 2. Carregamento dos datasets ─────────────────────────────────────────────

def carregar_dados():
    print("Carregando datasets...")
    areas_gdf = gpd.read_file("dados/bronze/sh_area_forca/areas_forca_municipal.shp").set_crs(CRS, allow_override=True)

    ocorr = pl.read_csv(
        "dados/bronze/df_ocorrencias_tratado - Extração 1 .csv",
        infer_schema_length=10000,
    ).filter(
        pl.col("longitude").is_between(-44.0, -43.0)
        & pl.col("latitude").is_between(-23.2, -22.7)
    )
    ocorr_pd = ocorr.to_pandas()
    ocorr_gdf = gpd.GeoDataFrame(
        ocorr_pd,
        geometry=gpd.points_from_xy(ocorr_pd["longitude"], ocorr_pd["latitude"]),
        crs=CRS,
    )

    fu = pl.read_csv("dados/bronze/fatores_urbanos.csv", infer_schema_length=0)
    dd = pd.read_csv("dados/bronze/disk_denuncia.csv", sep=";", encoding="latin-1", low_memory=False)
    cam = pl.read_csv("dados/bronze/cameras_areas_fm.csv", infer_schema_length=0)

    print(f"  {len(ocorr_gdf):,} ocorrências | {len(fu):,} fatores | {len(dd):,} denúncias | {len(cam):,} câmeras")
    return areas_gdf, ocorr_gdf, fu, dd, cam


# ── 3. Preparação dos dados por área ─────────────────────────────────────────

def _hora_bin(h) -> str | None:
    if pd.isna(h):
        return None
    try:
        hh = int(str(h).split(":")[0])
    except (ValueError, AttributeError):
        return None
    if hh < 6:
        return "madrugada (00h–06h)"
    if hh < 12:
        return "manhã (06h–12h)"
    if hh < 18:
        return "tarde (12h–18h)"
    return "noite (18h–24h)"


def preparar_dados(nome_subar, areas_gdf, ocorr_gdf, fu_df, dd_df, cam_df) -> dict | None:
    cfg = AREA_CONFIG[nome_subar]
    area_poly = areas_gdf[areas_gdf["nome_subar"] == nome_subar]

    # Ocorrências dentro do polígono da área FM
    ocorr_area = gpd.sjoin(ocorr_gdf, area_poly[["geometry"]], how="inner", predicate="within")
    if len(ocorr_area) == 0:
        return None

    delito_counts = ocorr_area["desc_delito"].value_counts().head(5).to_dict()
    top_locf = (
        ocorr_area["locf"].dropna().str.title().value_counts().head(8).to_dict()
    )
    dia_counts = ocorr_area["dia_semana"].dropna().value_counts().to_dict()
    hora_dist = dict(Counter(ocorr_area["hora"].apply(_hora_bin).dropna()))
    ano_counts = {int(k): int(v) for k, v in ocorr_area["ano"].value_counts().sort_index().items()}

    # Câmeras da área
    cam_area = cam_df.filter(pl.col("nome_area_fm") == cfg["cam_nome"])
    n_cameras = len(cam_area)

    # Fatores urbanos da área
    fu_area = fu_df.filter(pl.col("subarea_nome") == nome_subar)
    fat_counts = (
        fu_area
        .group_by(["tipo_ocorrencia_descricao", "orgao_responsavel"])
        .agg(pl.len().alias("n"))
        .sort("n", descending=True)
        .to_pandas()
        .to_dict(orient="records")
    )[:15]

    # Bairros derivados dos fatores → filtrar disk denúncia
    bairros = [b.upper() for b in fu_area["bairro_nome"].drop_nulls().unique().to_list()]

    dd_classes = [
        "CRIMES CONTRA O PATRIMÔNIO",
        "CRIMES CONTRA A PESSOA",
        "SUBSTÂNCIAS ENTORPECENTES",
        "PERTURBAÇÃO DA ORDEM PÚBLICA",
    ]
    dd_area = dd_df[
        dd_df["bairro_logradouro"].str.upper().isin(bairros)
        & dd_df["assuntos.classe"].isin(dd_classes)
    ]
    relatos = dd_area["relato_redacted"].dropna().head(8).tolist()

    return {
        "area_nome": cfg["titulo"],
        "n_total": len(ocorr_area),
        "delito_counts": delito_counts,
        "top_locf": top_locf,
        "dia_counts": dia_counts,
        "hora_dist": hora_dist,
        "ano_counts": ano_counts,
        "n_cameras": n_cameras,
        "fat_counts": fat_counts,
        "bairros": bairros,
        "relatos_denuncia": relatos,
    }


# ── 4. Chamada à API Claude ──────────────────────────────────────────────────

def chamar_claude(nome_subar: str, dados: dict, relints_info: dict, client: Anthropic) -> dict:
    cfg = AREA_CONFIG[nome_subar]
    ri = relints_info.get(nome_subar, {})
    subareas_ri = ri.get("subareas", [])
    fonte = ri.get("fonte", "N/A")

    subareas_txt = "\n".join(
        f'  Sub-área {i + 1}: "{s["nome"]}"\n'
        f'  Rota de fuga (fonte: {fonte}): {s["escape_route"] or "não extraída — gere com base na localização"}'
        for i, s in enumerate(subareas_ri)
    )

    delitos_txt = "\n".join(f"  {k}: {v}" for k, v in dados["delito_counts"].items())
    locf_txt = "\n".join(f"  {k}: {v}" for k, v in list(dados["top_locf"].items())[:5])
    fat_txt = "\n".join(
        f"  {r['tipo_ocorrencia_descricao']} ({r['orgao_responsavel']}): {r['n']}"
        for r in dados["fat_counts"]
    )
    hora_txt = " | ".join(
        f"{k}: {v}" for k, v in sorted(dados["hora_dist"].items())
    )
    relatos_txt = (
        "\n".join(f"  [{i + 1}] {r[:250]}" for i, r in enumerate(dados["relatos_denuncia"]))
        or "Nenhum relato disponível para os bairros desta área."
    )

    user_msg = f"""\
## ÁREA: {dados["area_nome"]}

### SUB-ÁREAS (extraídas do RELINT original — {fonte})
{subareas_txt}

### DADOS CRIMINAIS
- Total de ocorrências na área: {dados["n_total"]:,}
- Câmeras de vigilância: {dados["n_cameras"]}
- Tendência anual: {dados["ano_counts"]}
- Distribuição horária: {hora_txt}
- Dias mais críticos: {dados["dia_counts"]}

Tipos de delito (top 5):
{delitos_txt}

Logradouros com mais ocorrências:
{locf_txt}

### FATORES URBANOS (levantamento de campo — subárea "{nome_subar}")
{fat_txt}

### RELATOS DISQUE DENÚNCIA (bairros: {", ".join(dados["bairros"][:5])})
{relatos_txt}

---
Instruções:
- Use EXATAMENTE os nomes de sub-área listados acima.
- Para o fator 5 (rotas de dispersão), USE a rota do RELINT original quando disponível.
- Baseie a análise nos dados fornecidos (logradouros, horários, fatores, relatos).
- Retorne APENAS o JSON, sem blocos de código markdown ao redor.
"""

    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        system=[
            {
                "type": "text",
                "text": SYSTEM_PROMPT,
                "cache_control": {"type": "ephemeral"},
            }
        ],
        messages=[{"role": "user", "content": user_msg}],
    )

    raw = response.content[0].text.strip()
    # Strip accidental markdown fences
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1].lstrip("json").strip() if len(parts) > 1 else raw
    return json.loads(raw)


# ── 5. Montagem do .docx ─────────────────────────────────────────────────────

def _set_cell_content(cell, paragraphs_data: list[dict]) -> None:
    """Substitui todo o conteúdo da célula pelos parágrafos especificados.

    paragraphs_data: list of {"text": str, "bold": bool (optional)}
    """
    tc = cell._tc
    for p_elem in list(tc.findall(qn("w:p"))):
        tc.remove(p_elem)

    for pdata in paragraphs_data:
        p = OxmlElement("w:p")
        text = pdata.get("text", "")
        if text:
            r = OxmlElement("w:r")
            if pdata.get("bold"):
                rpr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rpr.append(b)
                r.append(rpr)
            t = OxmlElement("w:t")
            t.text = text
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r.append(t)
            p.append(r)
        tc.append(p)


def montar_docx(nome_subar: str, conteudo: dict, relints_info: dict) -> pathlib.Path:
    cfg = AREA_CONFIG[nome_subar]
    ri = relints_info[nome_subar]

    template_path = RELINTS_DIR / ri["fonte"]
    doc = Document(template_path)
    table = doc.tables[0]

    # Row 2: intro
    _set_cell_content(table.rows[2].cells[0], [{"text": conteudo["intro"]}])

    # Rows 3–8: sub-áreas (3 × (nome row + content row))
    subareas_ri = ri["subareas"]
    for i, (sa, row_nome_idx) in enumerate(zip(conteudo["subareas"], [3, 5, 7])):
        nome_original = subareas_ri[i]["nome"]

        # Nome da sub-área (bold)
        _set_cell_content(
            table.rows[row_nome_idx].cells[0],
            [{"text": nome_original, "bold": True}],
        )

        # Conteúdo da sub-área
        paras = [
            {"text": sa["abertura"]},
            {"text": ""},
            {"text": "Também foram identificados:"},
        ]
        for fator in sa["fatores"]:
            paras.append({"text": f"• {fator}"})
        paras += [
            {"text": ""},
            {"text": sa["dinamica"]},
        ]
        _set_cell_content(table.rows[row_nome_idx + 1].cells[0], paras)

    # Row 9: "CONCLUSÃO" — mantém do template

    # Row 10: conclusão
    concl = conteudo["conclusao"]
    paras_concl = [
        {"text": concl["abertura"]},
        {"text": ""},
    ]
    for rec in concl["recomendacoes"]:
        paras_concl.append({"text": f"• {rec}"})
    paras_concl += [
        {"text": ""},
        {"text": concl["fechamento"]},
    ]
    _set_cell_content(table.rows[10].cells[0], paras_concl)

    OUTPUT_DIR.mkdir(exist_ok=True)
    slug = (
        cfg["titulo"]
        .replace(" ", "_")
        .replace("–", "-")
        .replace("—", "-")
        .replace("/", "-")
    )
    out_path = OUTPUT_DIR / f"RI_{cfg['ri_num']}_2026_{slug}.docx"
    doc.save(out_path)
    return out_path


# ── 6. Orquestração principal ────────────────────────────────────────────────

def main() -> None:
    target = sys.argv[1] if len(sys.argv) > 1 else None

    if target:
        areas = [a for a in AREA_CONFIG if target.lower() in a.lower()]
        if not areas:
            print(f"Área '{target}' não encontrada. Disponíveis:")
            for a in AREA_CONFIG:
                print(f"  - {a}")
            sys.exit(1)
    else:
        areas = list(AREA_CONFIG.keys())

    print("1. Extraindo RELINTs existentes...")
    relints_info = extrair_relints()
    print(f"   {len(relints_info)} RELINTs carregados")

    print("\n2. Carregando datasets...")
    areas_gdf, ocorr_gdf, fu_df, dd_df, cam_df = carregar_dados()

    client = Anthropic()

    print(f"\n3. Gerando {len(areas)} relatório(s)...")
    for nome_subar in areas:
        cfg = AREA_CONFIG[nome_subar]
        print(f"\n  [{cfg['ri_num']}] {nome_subar}")

        if nome_subar not in relints_info:
            print("    AVISO: RELINT original não encontrado — pulando")
            continue

        dados = preparar_dados(nome_subar, areas_gdf, ocorr_gdf, fu_df, dd_df, cam_df)
        if not dados:
            print("    AVISO: sem ocorrências na área — pulando")
            continue

        print(
            f"    {dados['n_total']:,} ocorrências | {dados['n_cameras']} câmeras"
            f" | {len(dados['relatos_denuncia'])} relatos"
        )

        print("    Chamando Claude API...")
        conteudo = chamar_claude(nome_subar, dados, relints_info, client)

        out_path = montar_docx(nome_subar, conteudo, relints_info)
        print(f"    -> {out_path}")

    print("\nConcluído.")


if __name__ == "__main__":
    main()
